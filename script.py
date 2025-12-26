from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title section
titles = [
    "FORM ‘A’",
    "MEDIATION APPLICATION FORM",
    "[REFER RULE 3(1)]",
    "",
    "Mumbai District Legal Services Authority",
    "City Civil Court, Mumbai",
    ""
]

for t in titles:
    p = doc.add_paragraph(t)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if t.strip():
        p.runs[0].bold = True

# Main table
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"

def add_row(left, right):
    row = table.add_row().cells
    row[0].text = left
    row[1].text = right

add_row("DETAILS OF PARTIES:", "")
add_row("1 Name of Applicant", "{{client_name}}")
add_row("Address and contact details of Applicant", "")
add_row("Address",
        "REGISTERED ADDRESS:\n{{branch_address}}\n\n"
        "CORRESPONDENCE BRANCH ADDRESS:\n{{branch_address}}")
add_row("Telephone No.", "{{mobile}}")
add_row("Mobile No.", "")
add_row("Email ID", "info@kslegal.co.in")

add_row("2 Name, Address and Contact details of Opposite Party:", "")
add_row("Address and contact details of Defendant/s", "")
add_row("Name", "{{customer_name}}")
add_row(
    "Address",
    "REGISTERED ADDRESS:\n"
    "{% if address1 and address1 != '' %}{{address1}}{% else %} ________________ {% endif %}\n\n"
    "CORRESPONDENCE ADDRESS:\n"
    "{% if address1 and address1 != '' %}{{address1}}{% else %} ________________ {% endif %}"
)
add_row("Telephone No.", "")
add_row("Mobile No.", "")
add_row("Email ID", "")

add_row("DETAILS OF DISPUTE:", "")
add_row("", "THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES, 2018")
add_row(
    "",
    "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
)

doc.save("Mediation_Application_Form.docx")
